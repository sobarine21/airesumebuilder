import streamlit as st
import google.generativeai as genai
from fpdf import FPDF
from docx import Document
import os

# Configure the Gemini API key securely from Streamlit's secrets
genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])

# Streamlit App UI
st.title("AI-Powered Resume Builder")
st.write("Use AI to create a personalized and enhanced resume based on your input.")

# User input fields
name = st.text_input("Full Name")
contact = st.text_input("Contact Information (Email, Phone, etc.)")
objective = st.text_area("Career Objective")
education = st.text_area("Education Details (Degree, School, Year)")
experience = st.text_area("Work Experience (Company, Role, Years)")
skills = st.text_area("Skills (e.g., Python, Communication, Project Management)")
awards = st.text_area("Awards & Certifications (Optional)")

# Additional Sections (Optional)
hobbies = st.text_area("Hobbies and Interests (Optional)")
volunteer = st.text_area("Volunteer Experience (Optional)")
projects = st.text_area("Projects (Optional)")

# Choose Resume Template
template_choice = st.selectbox("Select Resume Template", ["Minimalist", "Professional", "Creative"])

# Option to upload a Profile Picture
profile_pic = st.file_uploader("Upload Profile Picture (Optional)", type=["jpg", "jpeg", "png"])

# Job Role for Resume Tailoring
job_role = st.text_input("Job Role (e.g., Software Engineer, Marketing Manager, etc.)")

# Button to generate AI-enhanced resume
if st.button("Generate Resume"):
    if name and contact and objective and education and experience and skills:
        try:
            # Prepare the prompt for Gemini AI based on inputs
            prompt = f"""
            Create a professional resume using the following details:

            Name: {name}
            Contact Info: {contact}
            Objective: {objective}
            Education: {education}
            Work Experience: {experience}
            Skills: {skills}
            Awards: {awards}
            Hobbies: {hobbies}
            Volunteer Experience: {volunteer}
            Projects: {projects}
            Job Role: {job_role}
            Template: {template_choice}

            Format it in a clean, ATS-friendly structure with proper headings. 
            Adjust the tone and style according to the selected template.
            """

            # Load Gemini AI model and generate the resume content
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(prompt)
            resume_text = response.text

            # Display the AI-generated resume in Streamlit
            st.subheader("AI-Generated Resume")
            st.write(resume_text)

            # Convert the generated resume text into a PDF
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            # Add profile picture if uploaded
            if profile_pic:
                img_path = f"/tmp/{name}_profile.jpg"
                with open(img_path, "wb") as f:
                    f.write(profile_pic.getbuffer())
                pdf.image(img_path, x=10, y=10, w=30)

            # Add resume content to the PDF
            pdf.set_font("Arial", size=12)
            pdf.ln(35)  # Add space after the image
            pdf.multi_cell(0, 10, resume_text)

            # Save the generated PDF to a file
            resume_pdf = "/tmp/generated_resume.pdf"
            pdf.output(resume_pdf)

            # Provide the download button for the PDF
            with open(resume_pdf, "rb") as f:
                st.download_button(
                    label="Download Your AI-Generated Resume (PDF)",
                    data=f,
                    file_name="AI_Resume.pdf",
                    mime="application/pdf"
                )

            # Convert the generated resume text into DOCX format
            doc = Document()
            doc.add_heading(f"Resume of {name}", 0)
            doc.add_paragraph(resume_text)

            # Save the DOCX file
            resume_docx = "/tmp/generated_resume.docx"
            doc.save(resume_docx)

            # Provide the download button for the DOCX
            with open(resume_docx, "rb") as f:
                st.download_button(
                    label="Download Your AI-Generated Resume (Word)",
                    data=f,
                    file_name="AI_Resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error(f"Error: {e}")
    else:
        st.warning("Please fill in all the fields before generating the resume.")

# Footer
st.markdown("""
---
*Powered by Gemini AI and Streamlit.*
""")
